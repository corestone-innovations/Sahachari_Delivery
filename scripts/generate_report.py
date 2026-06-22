from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import os

root = Path.cwd()
output_path = root / 'SAHACHARI_DELIVERY_REPORT.docx'
image_dir = root / 'screenshots'


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if size != 11:
            run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_image(doc, path, width=Inches(4.5)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=width)
        return True
    return False


def build_report():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run('Sahachari Delivery')
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = p2.add_run('Internship Project Report')
    r2.font.size = Pt(15)
    r2.bold = True

    doc.add_paragraph()
    p3 = doc.add_paragraph('Submitted by: Muhammed Yasir M S')
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p4 = doc.add_paragraph('Master of Vocation in Software Application Development')
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p5 = doc.add_paragraph('Cochin University of Science and Technology')
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    add_heading(doc, 'DECLARATION', 1)
    add_paragraph(doc, 'I hereby declare that this internship report entitled “Sahachari Delivery” is a record of my own work carried out during my internship and has not been submitted elsewhere for the award of any degree or diploma.')
    add_paragraph(doc, 'Name: Muhammed Yasir M S')
    add_paragraph(doc, 'Place: CUSAT')
    add_paragraph(doc, 'Date: June 2026')
    doc.add_page_break()

    add_heading(doc, 'ACKNOWLEDGMENT', 1)
    add_paragraph(doc, 'I would like to express my sincere gratitude to my guide, faculty members, and the team at Corestone Innovations for their guidance and support throughout this internship project.')
    add_paragraph(doc, 'I am also thankful to my friends and family for their encouragement and continuous support.')
    doc.add_page_break()

    add_heading(doc, 'ABSTRACT', 1)
    add_paragraph(doc, 'Sahachari Delivery is a mobile-based delivery management application developed to help delivery personnel manage their daily tasks efficiently. The application allows users to log in securely, view available jobs, accept orders, update status from accepted to picked up and delivered, collect payment, and access their profile. The project focuses on providing a simple, user-friendly workflow that improves delivery operations and makes order tracking easier.')
    doc.add_page_break()

    add_heading(doc, 'INTRODUCTION', 1)
    add_paragraph(doc, 'The Sahachari Delivery application is designed to support delivery personnel in handling deliveries in an organized and digital way. In traditional delivery operations, riders often rely on manual communication and paperwork, which can cause confusion and delays. This application provides a structured workflow where delivery tasks can be managed through a mobile interface.')
    add_paragraph(doc, 'The application helps users browse available jobs, accept orders, track order progress, collect payments, and maintain their profile details. The system is developed using React Native and Expo, providing a modern and smooth experience for users.')
    doc.add_page_break()

    add_heading(doc, 'ABOUT THE APPLICATION', 1)
    add_paragraph(doc, 'Sahachari Delivery is a mobile application focused on the delivery workflow of a delivery person. It allows the user to view available deliveries, accept a task, update order progress, collect payment, and manage personal details. The application is designed to be simple, practical, and suitable for daily use in the delivery process.')
    add_bullets(doc, [
        'Secure login for delivery personnel',
        'Available jobs page to browse orders',
        'Accept job feature to start a delivery',
        'Order tracking from accepted to delivered',
        'Payment collection after delivery completion',
        'Profile section for managing personal information'
    ])
    doc.add_page_break()

    add_heading(doc, 'OBJECTIVE AND PURPOSE', 1)
    add_paragraph(doc, 'The main objective of this project is to develop a delivery management application that makes the process faster, simpler, and more reliable for delivery personnel. The app aims to reduce manual effort, improve order visibility, and support the user in completing deliveries smoothly.')
    add_paragraph(doc, 'The purpose of the system is to provide an organized digital workflow for delivery tasks while ensuring that the user can track each order from receiving the job to completion and payment collection.')
    doc.add_page_break()

    add_heading(doc, 'SCOPE', 1)
    add_bullets(doc, [
        'Provide a login system for delivery personnel',
        'Display available delivery jobs',
        'Allow users to accept and manage assigned deliveries',
        'Track the status of deliveries',
        'Enable payment collection for completed orders',
        'Provide a profile page for managing account details'
    ])
    doc.add_page_break()

    add_heading(doc, 'REQUIREMENT SPECIFICATION', 1)
    add_paragraph(doc, 'The project requires a mobile application environment with a modern frontend framework, navigation, API integration, and local state management. The application is built using React Native with Expo, and it connects to backend services for user authentication, jobs, orders, and payments.')
    add_bullets(doc, [
        'Frontend: React Native, Expo, Expo Router',
        'State Management: React Query and React Context',
        'Authentication: JWT token-based login',
        'Backend Integration: REST APIs',
        'Platform: Android and mobile-supported interface'
    ])
    doc.add_page_break()

    add_heading(doc, 'SYSTEM ARCHITECTURE', 1)
    add_paragraph(doc, 'The Sahachari Delivery application follows a simple mobile architecture where the user interacts with the app interface, the app communicates with backend services, and order data is updated in the system. The architecture includes the login screen, jobs screen, order tracking screen, payment screen, and profile screen, all connected through a shared application flow.')
    add_paragraph(doc, 'Basic Data Flow Diagram')
    add_paragraph(doc, 'User → Login → Available Jobs → Accept Job → Order Tracking → Collect Payment → Profile')
    doc.add_page_break()

    add_heading(doc, 'SCREENSHOTS', 1)
    screenshots = [
        ('1.jpeg', 'Login Screen'),
        ('2.jpeg', 'Available Jobs Page'),
        ('3.jpeg', 'Order Accepted'),
        ('4.jpeg', 'Order Picked Up'),
        ('5.jpeg', 'Collect Payment'),
        ('6.jpeg', 'Order Delivered'),
        ('7.jpeg', 'Profile Screen'),
    ]
    for file_name, title in screenshots:
        add_heading(doc, title, 2)
        path = image_dir / file_name
        if add_image(doc, path):
            doc.add_paragraph()
        else:
            add_paragraph(doc, f'Screenshot not found: {file_name}')
    doc.add_page_break()

    add_heading(doc, 'CONCLUSION', 1)
    add_paragraph(doc, 'The Sahachari Delivery application is a practical and useful mobile solution for delivery personnel. It provides a clear and simple way to manage jobs, update status, collect payment, and view profile information. The project demonstrates the effective use of modern mobile development tools and provides a strong foundation for future improvement and expansion.')
    doc.add_page_break()

    add_heading(doc, 'FUTURE SCOPE', 1)
    add_bullets(doc, [
        'Add live order tracking and push notifications',
        'Improve payment integration and digital receipts',
        'Add admin dashboard for monitoring deliveries',
        'Support offline use and better data sync'
    ])

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    generated = build_report()
    print(f'Created {generated}')
